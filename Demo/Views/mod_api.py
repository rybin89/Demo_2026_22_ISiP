"""
Модуль для работы с API и валидации данных.

Предоставляет графический интерфейс для получения ФИО из внешнего API,
валидации полученных данных и сохранения результатов тестирования
в файл формата DOCX.
"""

import json
import re
from tkinter import *
from tkinter import ttk

from requests import get

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ModApi(Tk):
    """Графическое приложение для валидации данных ФИО.

    Представляет собой окно Tkinter, которое позволяет:
    - Получить ФИО из внешнего API
    - Проверить ФИО на соответствие допустимым символам (только кириллица)
    - Сохранить результат валидации в таблицу файла ТестКейс.docx

    Attributes:
        test_case_file (str): Путь к файлу с тестовыми кейсами (.docx).
        api_url (str): URL внешнего API для получения ФИО.
        fullname (str): Последнее полученное ФИО для валидации.
    """

    def __init__(self):
        """Инициализация главного окна приложения.

        Настраивает геометрию окна, создаёт элементы интерфейса:
        - Кнопку и метку для получения данных из API
        - Кнопку и метку для отправки результатов валидации
        """
        super().__init__()

        self.test_case_file = "ТестКейс.docx"

        # получить api
        self.api_url = 'http://prb.sylas.ru/TransferSimulator/fullName'
        # конфигурация окна
        self.title("Валидация данных")
        self.geometry("500x250")
        self.data_api = ttk.Frame(self)
        self.data_api.pack(anchor="center", fill = X, padx = 10, pady = 10)
        self.data_button = ttk.Button(self.data_api,text='Получить данные',command=self.get_fullname)
        self.data_button.grid(sticky=NSEW, row = 0, column =0, ipadx=50, ipady=6, padx=4, pady=4)
        self.data_fullname = ttk.Label(self.data_api, text='Text')
        self.data_fullname.grid(sticky=E,row = 0, column =1, ipadx=6, ipady=6, padx=4, pady=4)

        # Валидация
        self.validate = ttk.Frame(self)
        self.validate.pack(anchor="center", fill = X, padx = 10, pady = 10)
        self.validate_button = ttk.Button(self.validate, text='Отправить результат теста',command=lambda: self.validate_fullname_button(self.fullname) )
        self.validate_button.grid(sticky=NSEW, row=0, column=0, ipadx=50, ipady=6, padx=4, pady=4)
        self.validate_fullname = ttk.Label(self.validate)
        self.validate_fullname.grid(sticky=E, row=0, column=1, ipadx=6, ipady=6, padx=4, pady=4)
        self.fullname = ''

    def get_fullname(self):
        """Получает ФИО из внешнего API и отображает его в интерфейсе.

        Выполняет GET-запрос к API, извлекает значение 'value' из JSON-ответа
        и обновляет метку на форме, а также сохраняет значение в self.fullname.
        """
        response = get(self.api_url,{'key':'value'})

        print(response.json())
        fullname = response.json()['value']
        self.data_fullname['text'] = fullname
        self.fullname = fullname



    def validate_fullname_button(self, fullname):
        """Проверяет ФИО на соответствие допустимым символам.

        Выполняет валидацию ФИО с помощью регулярного выражения:
        строка должна содержать три слова из кириллических символов,
        разделённых пробелами (Фамилия Имя Отчество).

        Если валидация не пройдена, из строки удаляются запрещённые символы
        и результат сохраняется как "Не успешно".

        Args:
            fullname (str): ФИО для проверки.

        Returns:
            bool: True, если валидация прошла успешно, иначе False.
        """
        ## Проверка на допусимые символы
        #переменная для регулярного вырожения
        pattern = r'^[а-яА-ЯёЁ]+\s[а-яА-ЯёЁ]+\s[а-яА-ЯёЁ]+$'
        if re.fullmatch(pattern,fullname):
            self.validate_fullname['text'] = 'Валидация прошла успешно'
            self.save_test_result(fullname,"Успешно")
            return True
        else:
            self.validate_fullname['text'] = 'ФИО содержит запрещённые символы'
            fullname = self.clear_fullname(fullname,pattern)
            self.save_test_result(fullname, "Не успешно")
            return False

    def clear_fullname(self, fullname: str, pattern: str) -> str:
        """Удаляет из ФИО все символы, кроме кириллицы, пробелов и дефисов.

        Args:
            fullname (str): Исходная строка ФИО.
            pattern (str): Регулярное выражение для проверки (не используется в фильтрации).

        Returns:
            str: Очищенная строка, содержащая только допустимые символы.
        """
        return re.sub(r'[^а-яА-ЯёЁ\s-]', '', fullname)

    def save_test_result(self, fullname: str, result: str):
        """Сохраняет результат валидации в файл ТестКейс.docx.

        Добавляет новую строку в первую таблицу документа .docx
        с номером, ФИО и результатом проверки. Все ячейки форматируются:
        выравнивание по центру, размер шрифта 11pt.

        Args:
            fullname (str): ФИО для сохранения.
            result (str): Результат валидации ("Успешно" или "Не успешно").
        """
        try:

            doc = Document(self.test_case_file)
            row_cells = doc.tables[0].add_row().cells
            row_cells[0].text = str(len(doc.tables[0].rows) - 1)
            row_cells[1].text = fullname
            row_cells[2].text = result


        except:
            print("Ошибка")

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(11)


        doc.save(self.test_case_file)


if __name__ == "__main__":
    window = ModApi()
    window.mainloop()
