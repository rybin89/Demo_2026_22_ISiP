# -*- coding: utf-8 -*-
"""
Руководство по использованию Tkinter окон для управления записями БД
и капча-пазла для защиты от ботов.

Данный модуль содержит подробное описание того, как Tkinter-окна
(представления/View) проекта Demo взаимодействуют с Контроллером
и Моделью для выполнения CRUD-операций над таблицей user в базе данных,
а также описание модулей капча-пазла Puzzl_1 и API-валидатора ModApi.

---

## 1. Архитектура взаимодействия

В проекте Demo используется архитектура Model-View-Controller (MVC):

- **Model (Модель)** - ORM Peewee (User) - работа с базой данных
- **Controller (Контроллер)** - UserController - бизнес-логика
- **View (Представление)** - Tkinter окна (AuthView, AdminView, EditUser) - интерфейс пользователя

Схема взаимодействия:

```
Tkinter Окно (View)
    | вызывает методы
    v
UserController (Controller)
    | использует методы Peewee
    v
User (Model/ORM) ---> База данных MySQL
```

---

## 2. Окно авторизации (AuthView)

Модуль: ``Demo.Views.AuthView``

Окно авторизации - точка входа в приложение.
Позволяет пользователю войти в систему и, если он администратор,
перейти к панели управления пользователями.

### 2.1. Структура окна

```python
class AuthView(Tk):
    # Поля ввода
    login_entry = ttk.Entry()      # поле ввода логина
    password_entry = ttk.Entry()   # поле ввода пароля
    login_message = ttk.Label()    # метка для сообщений

    # Кнопка входа
    login_button = ttk.Button(text='Войти', command=self.login)
```

### 2.2. Обработка авторизации (READ)

Метод ``login()``:
1. Получает логин и пароль из полей ввода.
2. Вызывает ``UserController.auth(login, password)`` для проверки.
3. Если успешно и роль 'admin' - открывает ``AdminView`` (панель администратора).

```python
def login(self):
    login = self.login_entry.get()
    password = self.password_entry.get()
    user = UserController.auth(login, password)
    self.login_message['text'] = user['message']
    if user['success']:
        if user['user'].role == 'admin':
            window = AdminView()
            self.destroy()
```

---

## 3. Панель администратора (AdminView)

Модуль: ``Demo.Views.AdminView``

Панель администратора - основное окно для управления записями таблицы user.
Содержит два функциональных блока:

1. **Форма создания нового пользователя** (CREATE)
2. **Таблица всех пользователей** (READ + переход к UPDATE/DELETE)

### 3.1. Форма создания пользователя (CREATE)

Позволяет создать новую запись в таблице user:

```python
# Элементы формы
login_entry = ttk.Entry()          # поле ввода логина
password_entry = ttk.Entry()       # поле ввода пароля
combobox = ttk.Combobox()          # выбор роли (Пользователь/Администратор)
login_button = ttk.Button(text='Зарегистрировать', command=self.login)
```

Метод ``login()``:
1. Получает логин, пароль и роль из формы.
2. Преобразует роль из русского названия в системное значение.
3. Вызывает ``UserController.registration(login, password, role)``.
4. Обновляет таблицу пользователей.

```python
def login(self):
    login = self.login_entry.get()
    password = self.password_entry.get()
    role = self.roles_var.get()
    role = 'user' if role == 'Пользователь' else 'admin'
    if login == '' or password == '':
        self.login_message['text'] = 'Заполните все поля'
    else:
        user = UserController.registration(login, password, role)
        self.login_message['text'] = user['message']
    self.table()  # обновление таблицы
```

### 3.2. Таблица пользователей (READ)

Для отображения списка записей используется ``ttk.Treeview``:

```python
self.columns = ('login', 'role', 'blocked')
self.tree = ttk.Treeview(columns=self.columns, show="headings")
```

Метод ``table()``:
1. Очищает таблицу.
2. Получает список пользователей через ``UserController.get()``.
3. Преобразует данные в читаемый формат (роль и статус блокировки).
4. Заполняет таблицу строками.

```python
def table(self):
    for item in self.tree.get_children():
        self.tree.delete(item)

    users = UserController.get()['users']
    rows = []

    for user in users:
        user.bloked = 'Да' if user.bloked else 'Нет'
        user.role = 'Пользователь' if user.role == 'user' else 'Администратор'
        rows.append((user.login, user.role, user.bloked))

    self.tree.heading('login', text='Логин')
    self.tree.heading('role', text='Роль')
    self.tree.heading('blocked', text='Заблокирован')

    for item in rows:
        self.tree.insert('', END, values=item)
```

### 3.3. Выбор пользователя для редактирования

При выборе строки в таблице срабатывает обработчик ``item_selected()``,
который открывает окно редактирования:

```python
def item_selected(self, event):
    from Demo.Views.EditUser import EditUser
    self.item = self.tree.selection()[0]
    self.user_data = self.tree.item(self.item)['values'][0]  # логин
    window = EditUser(self.user_data)
    self.destroy()
```

---

## 4. Окно редактирования пользователя (EditUser)

Модуль: ``Demo.Views.EditUser``

Окно редактирования позволяет изменить данные существующего пользователя
(UPDATE). Принимает логин редактируемого пользователя и загружает
форму с полями для изменения.

### 4.1. Структура окна

```python
class EditUser(Tk):
    def __init__(self, login):
        self.login = login
        self.title(f'Изменение пользователя {self.login}')
        # Поля ввода
        login_entry = ttk.Entry()          # новый логин
        password_entry = ttk.Entry()       # новый пароль
        combobox = ttk.Combobox()          # новая роль
        blocked_button = ttk.Button()      # блокировка
        save_button = ttk.Button(text='Сохранить', command=self.save)
```

### 4.2. Сохранение изменений (UPDATE)

Метод ``save()``:
1. Получает новые значения из полей ввода.
2. Вызывает ``UserController.update()`` или ``UserController.new_update()``.
3. После сохранения возвращается к панели администратора.

```python
def save(self):
    login = self.login_entry.get()
    password = self.password_entry.get()
    role = self.roles_var.get()

    user = UserController.update(...)  # или new_update()
    self.login_message['text'] = user['message']
```

---

## 5. Сводная таблица: Окна и соответствующие CRUD-операции

| Tkinter окно    | CRUD операция | Метод контроллера    | Описание                         |
|-----------------|---------------|----------------------|----------------------------------|
| AuthView        | READ          | auth()               | Аутентификация пользователя      |
| AdminView       | CREATE        | registration()       | Создание нового пользователя     |
| AdminView       | READ          | get()                | Отображение списка пользователей |
| AdminView       | READ + UPDATE | item_selected()      | Выбор пользователя для редактирования |
| EditUser        | UPDATE        | update()/new_update()| Изменение данных пользователя    |
| EditUser        | DELETE        | delete()             | Удаление пользователя            |

---

## 6. Взаимодействие всех компонентов системы

### Полный цикл работы приложения:

1. **Запуск** - ``run.py`` создает ``AuthView``.
2. **Авторизация** - пользователь вводит логин/пароль.
3. **Проверка** - ``AuthView.login()`` вызывает ``UserController.auth()``.
4. **Панель администратора** - при успешном входе открывается ``AdminView``.
5. **Создание** - администратор заполняет форму и нажимает "Зарегистрировать".
6. **Обновление таблицы** - после создания вызывается ``AdminView.table()``.
7. **Редактирование** - выбор пользователя в таблице -> открытие ``EditUser``.
8. **Сохранение** - ``EditUser.save()`` -> ``UserController.update()``.

### Схема потока данных:

```
run.py
  |
  v
AuthView (авторизация)
  | UserController.auth()
  v (если admin)
AdminView
  |-- Форма (Create) --> UserController.registration() --> User.create()
  |-- Таблица (Read) --> UserController.get() --> User.select()
  |-- Выбор записи
        |
        v
      EditUser
          |-- Сохранение (Update) --> UserController.update() --> User.update()
          |-- Удаление (Delete) --> UserController.delete() --> User.delete()
```

---

## 7. Поля модели User, отображаемые в Tkinter окнах

| Поле модели | Тип      | Отображение в окнах           | Формат отображения              |
|-------------|----------|-------------------------------|---------------------------------|
| login       | str      | Entry (поле ввода)            | Текст                           |
| password    | str      | Entry (show='*')              | Скрытый текст                   |
| role        | str      | Combobox (выпадающий список)  | 'user' -> 'Пользователь'        |
|             |          |                               | 'admin' -> 'Администратор'      |
| bloked      | bool     | Button (Блокировать)          | False -> 'Нет' True -> 'Да'     |
| date_auth   | datetime | Не отображается               | -                               |

---

## 8. Капча-пазл (Puzzl_1) для защиты от ботов

Модуль: ``Demo.Views.Puzzl_1``

Реализует CAPTCHA-проверку в виде интерактивного пазла.
Пользователь должен собрать фрагменты изображения в правильном порядке
и ввести код "1234", чтобы подтвердить, что он человек.

### 8.1. Используемые библиотеки и модули

```python
from tkinter import *
from tkinter import messagebox
from PIL import Image, ImageTk
import random
```

- **PIL (Pillow)** - для обработки изображений (открытие, изменение размера).
- **random** - для случайного перемешивания фрагментов пазла.
- **messagebox** - для отображения сообщений пользователю.

### 8.2. Функция captcha_check()

Основная точка входа в модуль. Запускает окно капчи и возвращает
``True`` (капча пройдена) или ``False`` (капча не пройдена).

```python
def captcha_check():
    result = {'value': False}
    # ... отрисовка окна ...
    window.wait_window()
    return result['value']
```

Использование:
```python
from Demo.Views.Puzzl_1 import captcha_check

if captcha_check():
    print("Доступ разрешён")
else:
    print("Доступ запрещён")
```

### 8.3. Окно капчи

Создаёт главное окно Tk размером 450x650 с заголовком "Капча - Пазл":

```python
window = Tk()
window.geometry("450x650")
window.title("Капча - Пазл")
```

Элементы интерфейса:
- Заголовок "ПРОСТАЯ КАПЧА" (Label, font Arial 18 bold)
- Инструкция по переключению частей (Label, font Arial 10)
- Холст Canvas 300x300 для отрисовки пазла
- Поле ввода Entry для ввода порядка номеров
- Кнопка "Проверить" - вызов функции check()
- Кнопка "Перемешать" - вызов функции shuffle()

### 8.4. Изображения для пазла

Изображения хранятся в папке ``Demo/Views/images/`` в формате PNG:

| Файл    | Часть пазла        | Номер |
|---------|---------------------|-------|
| 1.png   | Верхняя-левая       | 1     |
| 2.png   | Верхняя-правая      | 2     |
| 3.png   | Нижняя-левая        | 3     |
| 4.png   | Нижняя-правая       | 4     |

Загрузка изображений через Pillow (PIL):

```python
for i in range(4):
    img = Image.open(f"images/{i + 1}.png")
    img = img.resize((140, 140), Image.Resampling.LANCZOS)
    pieces.append({
        'id': i,
        'correct': i,
        'pos': i,
        'img': ImageTk.PhotoImage(img)
    })
```

Если файл изображения не найден, создаётся серый прямоугольник-заглушка.

### 8.5. Перемешивание частей (shuffle)

Функция ``shuffle()`` случайным образом перемешивает позиции 4 частей:

```python
def shuffle():
    positions = list(range(4))
    random.shuffle(positions)
    for i, piece in enumerate(pieces):
        piece['pos'] = positions[i]
    show_puzzle()
```

Вызывается автоматически при старте и по нажатию кнопки "Перемешать".

### 8.6. Обработка кликов мыши (on_click)

Пользователь может менять две части местами, щёлкая по ним на холсте:

```python
def on_click(event):
    col = event.x // 150
    row = event.y // 150

    if selected_pos is None:
        # Первый клик - подсветить выбранную клетку
        selected_pos = clicked_pos
        canvas.create_rectangle(...)
    else:
        # Второй клик - поменять две части местами
        for p in pieces:
            if p['pos'] == selected_pos:
                p['pos'] = clicked_pos
            elif p['pos'] == clicked_pos:
                p['pos'] = selected_pos
        show_puzzle()
```

### 8.7. Отрисовка пазла (show_puzzle)

Рисует сетку 2x2 на холсте и размещает в каждой клетке часть изображения
с номером:

```python
def show_puzzle():
    canvas.delete("all")
    # Линии сетки
    for i in range(1, 3):
        canvas.create_line(0, i * 150, 300, i * 150, width=2)
        canvas.create_line(i * 150, 0, i * 150, 300, width=2)

    for piece in pieces:
        row = piece['pos'] // 2
        col = piece['pos'] % 2
        x, y = col * 150 + 5, row * 150 + 5
        canvas.create_image(x, y, image=piece['img'], anchor="nw")
        canvas.create_text(
            x + 70, y + 70, text=str(piece['id'] + 1),
            font=('Arial', 20, 'bold')
        )
```

### 8.8. Проверка правильности (check)

Сравнивает введённый пользователем порядок с эталоном "1234":

```python
def check():
    user_input = entry_order.get().strip()
    if user_input == "1234":
        result['value'] = True
        messagebox.showinfo("Успех!", "Капча пройдена!")
        window.destroy()
    else:
        messagebox.showerror("Ошибка!", f"Неправильно! Порядок: 1234")
```

### 8.9. Интеграция капчи с CRUD и Tkinter окнами

Капча может использоваться как дополнительный уровень защиты перед
выполнением критических операций. Например, перед редактированием
пользователя в ``AdminView.item_selected()``:

```python
def item_selected(self, event):
    from Demo.Views.Puzzl_1 import captcha_check
    if captcha_check():
        # Только после прохождения капчи - открыть окно редактирования
        from Demo.Views.EditUser import EditUser
        ...
```

---

## 9. Окно валидации данных из API (ModApi)

Модуль: ``Demo.Views.mod_api``

Предоставляет графический интерфейс для получения ФИО из внешнего REST API,
валидации полученных данных и сохранения результатов тестирования
в файл формата DOCX.

### 9.1. Используемые библиотеки и модули

```python
import json
import re
from tkinter import *
from tkinter import ttk

from requests import get

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
```

- **requests** — для выполнения HTTP-запроса к внешнему API.
- **re** — для валидации строк с помощью регулярных выражений.
- **docx (python-docx)** — для создания и редактирования DOCX-файлов.
- **tkinter.ttk** — для построения графического интерфейса.

### 9.2. Класс ModApi

Главное окно приложения, наследуется от ``Tk``.
Размер окна — 500x250 пикселей, заголовок — "Валидация данных".

```python
class ModApi(Tk):
    def __init__(self):
        super().__init__()
        self.title("Валидация данных")
        self.geometry("500x250")
        self.test_case_file = "ТестКейс.docx"
        self.api_url = 'http://prb.sylas.ru/TransferSimulator/fullName'
```

### 9.3. Атрибуты класса

| Атрибут          | Тип  | Описание                                    |
|------------------|------|---------------------------------------------|
| test_case_file   | str  | Путь к файлу ТестКейс.docx для записи       |
| api_url          | str  | URL внешнего API для получения ФИО          |
| fullname         | str  | Последнее полученное ФИО для валидации      |
| data_fullname    | Label| Метка для отображения полученного ФИО       |
| validate_fullname| Label| Метка для отображения результата валидации  |

### 9.4. Получение данных из API (get_fullname)

Метод выполняет GET-запрос к внешнему API, извлекает значение 'value'
из JSON-ответа и отображает его в интерфейсе:

```python
def get_fullname(self):
    response = get(self.api_url, {'key': 'value'})
    fullname = response.json()['value']
    self.data_fullname['text'] = fullname
    self.fullname = fullname
```

**Процесс:**
1. Отправляется GET-запрос на ``http://prb.sylas.ru/TransferSimulator/fullName``.
2. Сервер возвращает JSON вида ``{"value": "Иванов Иван Иванович"}``.
3. Значение извлекается из ответа и отображается в интерфейсе.

### 9.5. Валидация ФИО (validate_fullname_button)

Проверяет ФИО на соответствие регулярному выражению:
строка должна содержать три слова из кириллических символов,
разделённых пробелами (Фамилия Имя Отчество).

```python
def validate_fullname_button(self, fullname):
    pattern = r'^[а-яА-ЯёЁ]+\\s[а-яА-ЯёЁ]+\\s[а-яА-ЯёЁ]+$'
    if re.fullmatch(pattern, fullname):
        self.validate_fullname['text'] = 'Валидация прошла успешно'
        self.save_test_result(fullname, "Успешно")
        return True
    else:
        self.validate_fullname['text'] = 'ФИО содержит запрещённые символы'
        fullname = self.clear_fullname(fullname, pattern)
        self.save_test_result(fullname, "Не успешно")
        return False
```

**Алгоритм:**
1. Проверяет ФИО по регулярному выражению (только кириллица, три слова).
2. Если валидация пройдена — сохраняет результат "Успешно".
3. Если не пройдена — очищает строку от запрещённых символов и
   сохраняет результат "Не успешно".

### 9.6. Очистка строки от запрещённых символов (clear_fullname)

Удаляет из ФИО все символы, кроме кириллицы, пробелов и дефисов:

```python
def clear_fullname(self, fullname: str, pattern: str) -> str:
    return re.sub(r'[^а-яА-ЯёЁ\\s-]', '', fullname)
```

### 9.7. Сохранение результата в DOCX (save_test_result)

Сохраняет результат валидации в файл "ТестКейс.docx".
Добавляет новую строку в первую таблицу документа.

```python
def save_test_result(self, fullname: str, result: str):
    try:
        doc = Document(self.test_case_file)
        row_cells = doc.tables[0].add_row().cells
        row_cells[0].text = str(len(doc.tables[0].rows) - 1)
        row_cells[1].text = fullname
        row_cells[2].text = result
    except:
        print("Ошибка")

    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(11)

    doc.save(self.test_case_file)
```

**Структура таблицы в ТестКейс.docx:**
| Номер | ФИО                          | Результат   |
|-------|------------------------------|-------------|
| 1     | Иванов Иван Иванович          | Успешно     |
| 2     | Petrov Petr                   | Не успешно  |

### 9.8. Создание ModApi — пошаговая инструкция

**Шаг 1: Создание класса и настройка окна**

```python
class ModApi(Tk):
    def __init__(self):
        super().__init__()
        self.title("Валидация данных")
        self.geometry("500x250")
```

**Шаг 2: Определение внешних зависимостей**

```python
self.test_case_file = "ТестКейс.docx"
self.api_url = 'http://prb.sylas.ru/TransferSimulator/fullName'
```

**Шаг 3: Создание фрейма для API-запроса**

```python
self.data_api = ttk.Frame(self)
self.data_api.pack(anchor="center", fill=X, padx=10, pady=10)

self.data_button = ttk.Button(self.data_api, text='Получить данные',
                               command=self.get_fullname)
self.data_button.grid(sticky=NSEW, row=0, column=0, ipadx=50, ipady=6, padx=4, pady=4)

self.data_fullname = ttk.Label(self.data_api, text='Text')
self.data_fullname.grid(sticky=E, row=0, column=1, ipadx=6, ipady=6, padx=4, pady=4)
```

**Шаг 4: Создание фрейма для валидации**

```python
self.validate = ttk.Frame(self)
self.validate.pack(anchor="center", fill=X, padx=10, pady=10)

self.validate_button = ttk.Button(
    self.validate, text='Отправить результат теста',
    command=lambda: self.validate_fullname_button(self.fullname)
)
self.validate_button.grid(sticky=NSEW, row=0, column=0, ipadx=50, ipady=6, padx=4, pady=4)

self.validate_fullname = ttk.Label(self.validate)
self.validate_fullname.grid(sticky=E, row=0, column=1, ipadx=6, ipady=6, padx=4, pady=4)
```

**Шаг 5: Реализация методов CRUD-подобных операений**

| Операция  | Метод ModApi          | Действие                              |
|-----------|-----------------------|---------------------------------------|
| CREATE    | __init__              | Создание окна                         |
| READ (API)| get_fullname()        | Получение ФИО из внешнего API         |
| UPDATE    | validate_fullname_button() | Проверка и запись результата     |
| CREATE    | save_test_result()    | Добавление строки в DOCX-файл         |
| DELETE    | clear_fullname()      | Удаление запрещённых символов         |

---

## 10. Сводная таблица всех модулей Views

| Модуль          | Файл           | Тип                     | Связанные CRUD операции |
|-----------------|----------------|-------------------------|------------------------|
| AuthView        | AuthView.py    | Tk (TopLevel)           | READ (auth)            |
| AdminView       | AdminView.py   | Tk (TopLevel)           | CREATE + READ          |
| EditUser        | EditUser.py    | Tk (TopLevel)           | UPDATE + DELETE        |
| Puzzl_1         | Puzzl_1.py     | Функция + Tk (Toplevel) | Защита (CAPTCHA)       |
| ModApi          | mod_api.py     | Tk (TopLevel)           | API + валидация        |
| main            | main.py        | Точка входа             | Запуск приложения      |

---
"""

class AuthViewExample:
    """
    Пример окна авторизации (READ - чтение/проверка записи).

    Использует UserController.auth() для аутентификации пользователя.
    При успешном входе администратора открывает AdminView.

    Attributes:
        login_entry: Поле ввода логина.
        password_entry: Поле ввода пароля (символы скрыты).
        login_message: Метка для вывода сообщений.
    """
    login_entry: str = "admin"
    password_entry: str = "****"
    login_message: str = ""


class AdminViewExample:
    """
    Пример панели администратора (CREATE + READ).

    Содержит форму создания нового пользователя и таблицу
    со списком всех пользователей. Позволяет перейти к
    редактированию выбранного пользователя.

    Attributes:
        login_entry: Поле ввода логина нового пользователя.
        password_entry: Поле ввода пароля нового пользователя.
        roles_var: Выбор роли (Пользователь/Администратор).
        tree: Таблица Treeview с колонками: login, role, blocked.
    """
    login_entry: str = "new_user"
    password_entry: str = "****"
    roles_var: str = "Пользователь"


class EditUserExample:
    """
    Пример окна редактирования пользователя (UPDATE).

    Позволяет изменить логин, пароль, роль и статус блокировки
    выбранного пользователя.

    Attributes:
        login: Логин редактируемого пользователя.
        login_entry: Поле ввода нового логина.
        password_entry: Поле ввода нового пароля.
        roles_var: Выбор новой роли.
    """
    login: str = "admin"
    login_entry: str = "new_login"
    password_entry: str = "****"
    roles_var: str = "Администратор"


if __name__ == "__main__":
    print("Модуль-руководство по Tkinter окнам для управления записями БД")
    print("Запустите pdoc для генерации HTML-документации")
